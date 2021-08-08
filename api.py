from io import BytesIO
from model.database import Mysql
from mysql.connector import IntegrityError
from docx import Document
from config import mysql_conf, SUCCESS, ERROR, email_sender, BaseConfig
import smtplib
import email.message
from itsdangerous import TimedJSONWebSignatureSerializer, SignatureExpired


def insert(form_data):
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            INSERT INTO form (
            `type`, stu_name, `stu_id`, `stu_tel`, edu_sys, campus, department, grade, class, `stu_type`, 
            `stu_type_school`, `stu_type_department`, `stu_type_other`, pwd) 
            VALUES(
            '{form_data['type']}', '{form_data['stu-name']}', '{form_data['stu-id']}', '{form_data['stu-tel']}', 
            '{form_data['edu-sys']}', '{form_data['campus']}', '{form_data['department']}', '{form_data['grade']}', 
            '{form_data['class']}', '{form_data['stu-type']}', '{form_data['stu-type-school']}',
            '{form_data['stu-type-department']}', '{form_data['stu-type-other']}', '{form_data['pwd']}');
        '''
        try:
            cursor.execute(sql)
            db.commit()
        except:
            return ERROR, None

        sql = '''
            SELECT LAST_INSERT_ID();
        '''
        cursor.execute(sql)
        form_id = cursor.fetchone()['LAST_INSERT_ID()']

        try:
            for i in range(1, 7):
                sql = f'''
                    INSERT INTO course_transform (
                        form_id, `index`, get_term, get_name, get_credit, get_score, set_term, set_name, `set_credit`,
                        set_type)
                    VALUES(
                        '{form_id}', '{form_data['sbj-' + str(i) + '-index']}', '{form_data['sbj-' + str(i) + '-getTerm']}',
                        '{form_data['sbj-' + str(i) + '-getName']}', '{form_data['sbj-' + str(i) + '-getCredit']}',
                        '{form_data['sbj-' + str(i) + '-getScore']}', '{form_data['sbj-' + str(i) + '-setTerm']}',
                        '{form_data['sbj-' + str(i) + '-setName']}', '{form_data['sbj-' + str(i) + '-setCredit']}',
                        '{form_data['sbj-' + str(i) + '-setType']}');
                '''
                cursor.execute(sql)
            db.commit()
        except:
            return ERROR, None
    return SUCCESS, form_id


def insert_file(form_id):
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            UPDATE `credit-transfer`.form 

            SET file_name='{form_id}.pdf' 
            
            WHERE id={form_id};
        '''
        cursor.execute(sql)
        db.commit()
    return SUCCESS


def send_mail(**kw):
    msg = email.message.EmailMessage()
    msg['From'] = f"課程抵免申請系統<{email_sender['sender_account']}>"
    msg['To'] = kw['target_account']
    msg['Subject'] = kw['msg_subject']
    msg.add_attachment(kw['msg_content_html'], subtype='html')

    with smtplib.SMTP(host='smtp.gmail.com', port=587) as smtp:
        smtp.ehlo()
        smtp.starttls()
        smtp.login(email_sender['sender_account'], email_sender['sender_password'])
        smtp.send_message(msg)


def login(form_data):
    form_id = form_data['form-id']
    pwd = form_data['pwd']
    try:
        form_id = int(form_id)
    except:
        return ERROR
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            SELECT pwd FROM form WHERE id={form_id};
        '''
        cursor.execute(sql)
        results = cursor.fetchone()

    if results is None:
        return ERROR
    if pwd == results['pwd']:
        return SUCCESS
    else:
        return ERROR


def editable(form_data):
    form_id = form_data['form-id']
    try:
        form_id = int(form_id)
    except:
        return ERROR
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            SELECT `editable` FROM form WHERE id={form_id};
        '''
        cursor.execute(sql)
        results = cursor.fetchone()

    if results is None:
        return ERROR
    if results['editable'] == 1:
        return SUCCESS
    else:
        return ERROR


def fetch_form(form_data):
    form_id = form_data['form-id']
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            SELECT `id`,`type`, stu_name, `stu_id`, `stu_tel`, edu_sys, campus, department, grade, class, `stu_type`,
                `stu_type_school`, `stu_type_department`, `stu_type_other`, `file_name`, `editable`
            FROM `credit-transfer`.form 
            WHERE id={form_id};
        '''
        cursor.execute(sql)
        info = cursor.fetchone()

        sql = f'''
            SELECT get_term, get_name, get_credit, get_score, set_term, set_name, set_credit, set_type, set_verify,
                `index`
            FROM `credit-transfer`.course_transform 
            WHERE form_id={form_id};
        '''
        cursor.execute(sql)
        trans = cursor.fetchall()

    return SUCCESS, info, trans


def update_save(form_data, form_id):
    status_code = editable({'form-id': form_id})
    if status_code == ERROR:
        return ERROR
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            UPDATE `credit-transfer`.form 
            SET `type`='{form_data['type']}', stu_name='{form_data['stu-name']}', stu_id='{form_data['stu-id']}',
                stu_tel='{form_data['stu-tel']}', edu_sys='{form_data['edu-sys']}', campus='{form_data['campus']}', 
                department='{form_data['department']}', grade='{form_data['grade']}', class='{form_data['class']}',
                stu_type='{form_data['stu-type']}', stu_type_school='{form_data['stu-type-school']}', 
                stu_type_department='{form_data['stu-type-department']}', stu_type_other='{form_data['stu-type-other']}'
            WHERE id={form_id};
        '''
        try:
            cursor.execute(sql)
            db.commit()
        except:
            return ERROR
        try:
            for i in range(1, 7):
                sql = f'''
                    UPDATE `credit-transfer`.course_transform 
                    SET get_term='{form_data['sbj-' + str(i) + '-getTerm']}',
                        get_name='{form_data['sbj-' + str(i) + '-getName']}',
                        get_credit='{form_data['sbj-' + str(i) + '-getCredit']}',
                        get_score='{form_data['sbj-' + str(i) + '-getScore']}',
                        set_term='{form_data['sbj-' + str(i) + '-setTerm']}', 
                        set_name='{form_data['sbj-' + str(i) + '-setName']}',
                        set_credit='{form_data['sbj-' + str(i) + '-setCredit']}',
                        set_type='{form_data['sbj-' + str(i) + '-setType']}'
                    WHERE form_id={form_id} AND `index`='{i}';
                '''
                cursor.execute(sql)
            db.commit()
        except:
            return ERROR
    return SUCCESS


def delete_form(form_data):
    form_id = form_data['form-id']
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        for i in range(1, 7):
            sql = f'''
                DELETE FROM `credit-transfer`.course_transform WHERE form_id={form_id} AND `index`='{i}';
            '''
            cursor.execute(sql)
        db.commit()

        sql = f'''
            DELETE FROM `credit-transfer`.form WHERE id={form_id};
        '''
        cursor.execute(sql)
        db.commit()

    return SUCCESS


def admin_login_verfiy(form_data):
    username = form_data['username']
    pwd = form_data['pwd']
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            SELECT username, password 
            FROM `credit-transfer`.admin 
            WHERE username='{username}' AND password='{pwd}';
        '''
        cursor.execute(sql)
        result = cursor.fetchone()
    if result is None:
        return ERROR
    else:
        return SUCCESS


def admin_fetch_list():
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            SELECT id, `type`, stu_name, stu_id, stu_tel, edu_sys, campus, department, grade, class, stu_type, 
            pwd, file_name, editable 
            FROM `credit-transfer`.form
        '''
        cursor.execute(sql)
        results = cursor.fetchall()
    return SUCCESS, results


def admin_verify_submit(form_data, form_id):
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        for i in range(1, 7):
            sql = f'''
                UPDATE `credit-transfer`.course_transform 
                SET set_verify='{form_data['sbj-' + str(i) + '-setVerify']}'
                WHERE form_id={form_id} AND `index`='{i}';
            '''
            cursor.execute(sql)
        db.commit()
        sql = f'''
            UPDATE `credit-transfer`.form 
            SET editable=0
            WHERE id={form_id};
        '''
        cursor.execute(sql)
        db.commit()
    return SUCCESS


def admin_info_submit(form_data):
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            UPDATE `credit-transfer`.admin
            SET email='{form_data['email-address']}', real_name='{form_data['username']}'
            WHERE id=1;
        '''
        cursor.execute(sql)
        db.commit()
    return SUCCESS


def fetch_admin_info():
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = '''
            SELECT email, real_name FROM `credit-transfer`.admin WHERE id=1;
        '''
        cursor.execute(sql)
        results = cursor.fetchone()
    return SUCCESS, results


def form_download(db_info, db_trans, db_admin):
    def generate_docx(template_file_path, dictionary):
        temp = BytesIO()
        template_document = Document(template_file_path)

        for variable_key, variable_value in dictionary.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(temp)
        file = temp.getvalue()
        temp.close()
        return file

    def replace_text_in_paragraph(paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)

    def checkbox(text, target):
        try:
            target = int(target)
        except:
            return text
        index = 0
        for key, value in enumerate(text):
            if value == '□':
                index += 1
                if index == target:
                    text = text[:key] + '■' + text[key + 1:]
                    break
        return text

    def type_process_b(target):
        target = str(target)
        if target == '1':
            text = '''■共同必修□通識科目□專業科目
  ○必修
  ○選修
    '''
        elif target == '2':
            text = '''□共同必修■通識科目□專業科目
  ○必修
  ○選修
    '''
        elif target == '3':
            text = '''□共同必修□通識科目■專業科目
  ●必修
  ○選修
    '''
        elif target == '4':
            text = '''□共同必修□通識科目■專業科目
  ○必修
  ●選修
    '''
        else:
            text = '''□共同必修□通識科目□專業科目
  ○必修
  ○選修
    '''
        return text

    def type_process_a(target):
        target = str(target)
        if target == '1':
            text = '''■共同必修□通識科目□專業科目
  ○必修  ○選修
  ○外系選修
    '''
        elif target == '2':
            text = '''□共同必修■通識科目□專業科目
  ○必修  ○選修
  ○外系選修
    '''
        elif target == '3':
            text = '''□共同必修□通識科目■專業科目
  ●必修  ○選修
  ○外系選修
    '''
        elif target == '4':
            text = '''□共同必修□通識科目■專業科目
  ○必修  ●選修
  ○外系選修
    '''
        elif target == '5':
            text = '''□共同必修□通識科目■專業科目
  ○必修  ○選修
  ●外系選修
    '''
        else:
            text = '''□共同必修□通識科目□專業科目
  ○必修  ○選修
  ○外系選修
    '''
        return text

    stu_name = db_info['stu_name']
    stu_id = db_info['stu_id']
    stu_tel = db_info['stu_tel']
    buffer = '''□五專□二技□四技□碩士班□博士班'''
    edu_sys = checkbox(buffer, db_info['edu_sys'])
    buffer = '''□建工校區 □燕巢校區 □第一校區
□楠梓校區 □旗津校區'''
    campus = checkbox(buffer, db_info['campus'])
    department = db_info['department']
    grade = db_info['grade']
    class_id = db_info['class']
    if db_info['stu_type_school'] == '':
        db_info['stu_type_school'] = '____________________________ '
    if db_info['stu_type_department'] == '':
        db_info['stu_type_department'] = '____________________'
    if db_info['stu_type_other'] == '':
        db_info['stu_type_other'] = '_____________________________________________________'
    buffer_a = f'''學生身分別：□本年度新生(含本校重考入學新生)               □交換生
            □轉學生，原學校：{db_info['stu_type_school']} 原科系所別：{db_info['stu_type_department']}
            □學碩士一貫學生  □其他：{db_info['stu_type_other']}'''
    buffer_b = f'''學生身分別：□本校轉系生，原系科別：{db_info['stu_type_department']}   □復學生
            □新舊課程交替重補修學生，科目名稱與科目內容不同但性質相近者，需互為抵充。
            □其他：{db_info['stu_type_other']}'''
    if db_info['type'] == 'a':
        stu_type = checkbox(buffer_a, db_info['stu_type'])
    if db_info['type'] == 'b':
        stu_type = checkbox(buffer_b, db_info['stu_type'])

    variables = dict()
    variables['$(stu_name)'] = stu_name
    variables['$(stu_id)'] = stu_id
    variables['$(stu_tel)'] = stu_tel
    variables['$(edu_sys)'] = edu_sys
    variables['$(campus)'] = campus
    variables['$(department)'] = department
    variables['$(grade)'] = grade
    variables['$(class_id)'] = class_id
    variables['$(stu_type)'] = stu_type

    for i in range(0, 6):
        get_term = db_trans[i]['get_term']
        get_course = db_trans[i]['get_name']
        get_credit = db_trans[i]['get_credit']
        get_score = db_trans[i]['get_score']
        set_term = db_trans[i]['set_term']
        set_course = db_trans[i]['set_name']
        set_credit = db_trans[i]['set_credit']
        if db_info['type'] == 'a':
            set_type = type_process_a(db_trans[i]['set_type'])
        if db_info['type'] == 'b':
            set_type = type_process_b(db_trans[i]['set_type'])
        if db_trans[i]['set_verify'] == '0':
            admin_buffer = ''
        else:
            admin_buffer = db_admin['real_name']
        buffer = f'''□同意□不同意
審核人:
'''
        set_verify = checkbox(buffer, db_trans[i]['set_verify'])

        variables[f'$(get_term_{i + 1})'] = get_term
        variables[f'$(get_course_{i + 1})'] = get_course
        variables[f'$(get_credit_{i + 1})'] = get_credit
        variables[f'$(get_score_{i + 1})'] = get_score
        variables[f'$(set_term_{i + 1})'] = set_term
        variables[f'$(set_course_{i + 1})'] = set_course
        variables[f'$(set_credit_{i + 1})'] = set_credit
        variables[f'$(set_type_{i + 1})'] = set_type
        variables[f'$(set_verify_{i + 1})'] = set_verify
    if db_info['type'] == 'a':
        return SUCCESS, generate_docx('word/template-a.docx', variables)
    elif db_info['type'] == 'b':
        return SUCCESS, generate_docx('word/template-b.docx', variables)


def api_get_stu_id(stu_id):
    try:
        stu_id = int(stu_id)
    except:
        return ERROR
    with Mysql(mysql_conf) as db:
        cursor = db.cursor(dictionary=True)
        sql = f'''
            SELECT id, `type`, stu_name, stu_id, stu_tel, edu_sys, campus, department, grade, class, stu_type,
            stu_type_school, stu_type_department, stu_type_other, pwd, file_name, editable 
            FROM `credit-transfer`.form 
            WHERE stu_id={stu_id};
        '''
        cursor.execute(sql)
        info = cursor.fetchall()

        form_id_list = []
        for key, value in enumerate(info):
            form_id_list.append(value['id'])

        course_verify_list = []
        for form_id in form_id_list:
            sql = f'''
                SELECT get_term, get_name, get_credit, get_score, set_term, set_name, set_credit, set_type, set_verify,
                    `index`
                FROM `credit-transfer`.course_transform
                WHERE form_id={form_id};
            '''
            cursor.execute(sql)
            trans = cursor.fetchall()
            course_verify_list.append(trans)

    return course_verify_list


def registered_user(form_data):
    with Mysql(mysql_conf) as db:
        cur = db.cursor(dictionary=True)
        sql = f'''
            INSERT INTO `credit-transfer`.user(`mail`, `password`)
            VALUE (%(mail)s, %(password)s);
        '''
        bind = {
            'mail': form_data['mail'],
            'password': form_data['pwd']
        }
        try:
            cur.execute(sql, bind)
        except IntegrityError:
            sql = f'''
                        SELECT `verify` FROM `credit-transfer`.`user` WHERE mail='{form_data['mail']}'
                    '''
            cur.execute(sql)
            verify_state = cur.fetchone()['verify']
            if not verify_state:
                return 'Send verify mail'
            elif verify_state:
                return 'Already have this mail'

        db.commit()

    return SUCCESS


def generate_verify_token(form_data):
    signature = TimedJSONWebSignatureSerializer(BaseConfig().SECRET_KEY, expires_in=600, salt='email')
    token = signature.dumps({'mail': form_data['mail']})
    return token


def token_verify(token):
    signature = TimedJSONWebSignatureSerializer(BaseConfig().SECRET_KEY, salt='email')
    try:
        data = signature.loads(token)
    except SignatureExpired:
        return 'token expired'
    mail = dict(data).get('mail')
    if mail is None and mail == '':
        return ERROR
    with Mysql(mysql_conf) as db:
        cur = db.cursor(dictionary=True)
        sql = f'''
            UPDATE `credit-transfer`.`user` SET verify=1 WHERE mail='{mail}';
        '''
        cur.execute(sql)
        db.commit()
    return SUCCESS


def user_login(form_data):
    with Mysql(mysql_conf) as db:
        cur = db.cursor(dictionary=True)
        sql = f'''
            select mail, password, verify from `credit-transfer`.`user` where mail=%(mail)s; 
        '''
        bind = {
            'mail': form_data['mail']
        }
        try:
            cur.execute(sql, bind)
        except Exception as e:
            print(e)
            return ERROR
        else:
            results = cur.fetchone()
            if results is None:
                return ERROR
            elif results['password'] == form_data['pwd']:
                if results['verify'] == 1:
                    return SUCCESS
                elif results['verify'] == 0:
                    return ERROR
            else:
                return ERROR
